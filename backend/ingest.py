import pdfplumber
import chromadb
from chromadb.utils import embedding_functions
from sentence_transformers import SentenceTransformer
from config import *
from pathlib import Path
import re
from typing import List, Dict
import uuid
import pytesseract
from pdf2image import convert_from_path
import camelot
from datetime import datetime
import docx
from docx import Document
import extract_msg
import email
from email import policy
from email.parser import BytesParser

class DocumentProcessor:
    def __init__(self):
        self.embedding_model = SentenceTransformer(str(EMBEDDING_MODEL))
        self.client = chromadb.PersistentClient(path=str(CHROMA_DIR))
        self.collection = self.client.get_or_create_collection(
            name=COLLECTION_NAME,
            embedding_function=embedding_functions.SentenceTransformerEmbeddingFunction(
                model_name=str(EMBEDDING_MODEL)
            )  # type: ignore
        )

    def clean_text(self, text: str) -> str:
        return re.sub(r'\s+', ' ', text).strip()

    def extract_section_headers(self, text: str) -> List[str]:
        headers = re.findall(r'^(?:[A-Z][A-Z\s]+|\d+\.\s+.+|[A-Z]\..+)$', text, re.MULTILINE)
        return headers

    def extract_effective_dates(self, text: str) -> List[str]:
        date_patterns = [r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', r'\b\d{4}-\d{2}-\d{2}\b']
        dates = []
        for pat in date_patterns:
            dates += re.findall(pat, text)
        return dates

    def extract_footnotes(self, text: str) -> List[str]:
        footnotes = re.findall(r'^(\*|\d+)\s+.+$', text, re.MULTILINE)
        return footnotes

    def extract_tables(self, pdf_path: Path) -> List[str]:
        tables = []
        try:
            c_tables = camelot.read_pdf(str(pdf_path), pages='all', flavor='stream')  # type: ignore
            for table in c_tables:
                tables.append(table.df.to_string())
        except Exception:
            pass
        return tables

    def extract_tables_docx(self, doc: Document) -> List[str]:
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(' | '.join(cell.text.strip() for cell in row.cells))
            tables.append('\n'.join(rows))
        return tables

    def ocr_pdf(self, pdf_path: Path) -> str:
        try:
            images = convert_from_path(str(pdf_path))
            ocr_text = []
            for img in images:
                ocr_text.append(pytesseract.image_to_string(img))
            return '\n'.join(ocr_text)
        except Exception:
            return ''

    def semantic_chunking(self, text: str) -> List[str]:
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = []
        current_len = 0
        for sent in sentences:
            sent_len = len(sent.split())
            if current_len + sent_len > CHUNK_SIZE and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = current_chunk[-CHUNK_OVERLAP:]
                current_len = sum(len(s.split()) for s in current_chunk)
            current_chunk.append(sent)
            current_len += sent_len
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        return [chunk for chunk in chunks if len(chunk.split()) >= MIN_CHUNK_LENGTH]

    def store_chunks(self, chunks, metadatas):
        embeddings = self.embedding_model.encode(chunks).tolist()
        self.collection.add(
            documents=chunks,
            embeddings=embeddings,
            metadatas=metadatas,  # type: ignore
            ids=[str(uuid.uuid4()) for _ in chunks]
        )

    def process_pdf(self, pdf_path: Path):
        with pdfplumber.open(pdf_path) as pdf:
            full_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        if len(full_text.strip()) < 100:
            full_text = self.ocr_pdf(pdf_path)
        cleaned = self.clean_text(full_text)
        chunks = self.semantic_chunking(cleaned)
        if not chunks:
            return None
        section_headers = self.extract_section_headers(cleaned)
        effective_dates = self.extract_effective_dates(cleaned)
        footnotes = self.extract_footnotes(cleaned)
        tables = self.extract_tables(pdf_path)
        def safe_join(val):
            if isinstance(val, list):
                return "; ".join(str(v) for v in val) if val else ""
            return str(val) if val is not None else ""
        metadatas = [{
            "source": str(pdf_path.name),
            "page": "aggregated",
            "chunk_id": str(uuid.uuid4())[:8],
            "section_headers": safe_join(section_headers),
            "effective_dates": safe_join(effective_dates),
            "footnotes": safe_join(footnotes),
            "tables": safe_join(tables)
        } for _ in chunks]
        self.store_chunks(chunks, metadatas)
        return len(chunks)

    def process_docx(self, docx_path: Path):
        doc = Document(str(docx_path))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        cleaned = self.clean_text(full_text)
        chunks = self.semantic_chunking(cleaned)
        if not chunks:
            return None
        section_headers = self.extract_section_headers(cleaned)
        effective_dates = self.extract_effective_dates(cleaned)
        footnotes = self.extract_footnotes(cleaned)
        tables = self.extract_tables_docx(doc)
        def safe_join(val):
            if isinstance(val, list):
                return "; ".join(str(v) for v in val) if val else ""
            return str(val) if val is not None else ""
        metadatas = [{
            "source": str(docx_path.name),
            "page": "aggregated",
            "chunk_id": str(uuid.uuid4())[:8],
            "section_headers": safe_join(section_headers),
            "effective_dates": safe_join(effective_dates),
            "footnotes": safe_join(footnotes),
            "tables": safe_join(tables)
        } for _ in chunks]
        self.store_chunks(chunks, metadatas)
        return len(chunks)

    def process_email(self, email_path: Path):
        text = ""
        tables = []
        if email_path.suffix.lower() == ".msg":
            msg = extract_msg.Message(str(email_path))
            text = msg.body or ""
        elif email_path.suffix.lower() == ".eml":
            with open(email_path, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)  # type: ignore
                # Try to get plain text body
                try:
                    text = msg.get_body(preferencelist=('plain', 'html')).get_content()  # type: ignore
                except Exception:
                    # Fallback: join all payloads
                    if msg.is_multipart():
                        text = "\n".join(part.get_payload(decode=True).decode(errors='ignore')  # type: ignore
                                             for part in msg.walk() if part.get_content_type() == 'text/plain')
                    else:
                        text = msg.get_payload(decode=True).decode(errors='ignore') if msg.get_payload() else ""  # type: ignore
        cleaned = self.clean_text(text)
        chunks = self.semantic_chunking(cleaned)
        if not chunks:
            return None
        section_headers = self.extract_section_headers(cleaned)
        effective_dates = self.extract_effective_dates(cleaned)
        footnotes = self.extract_footnotes(cleaned)
        def safe_join(val):
            if isinstance(val, list):
                return "; ".join(str(v) for v in val) if val else ""
            return str(val) if val is not None else ""
        metadatas = [{
            "source": str(email_path.name),
            "page": "aggregated",
            "chunk_id": str(uuid.uuid4())[:8],
            "section_headers": safe_join(section_headers),
            "effective_dates": safe_join(effective_dates),
            "footnotes": safe_join(footnotes),
            "tables": safe_join(tables)
        } for _ in chunks]
        self.store_chunks(chunks, metadatas)
        return len(chunks)

def main():
    processor = DocumentProcessor()
    data_dir = Path(__file__).parent.parent / "data"
    for file_path in data_dir.glob("*.*"):
        if file_path.suffix.lower() == ".pdf":
            print(f"Processing {file_path.name} (PDF)...")
            chunk_count = processor.process_pdf(file_path)
            print(f"Added {chunk_count} chunks")
        elif file_path.suffix.lower() == ".docx":
            print(f"Processing {file_path.name} (Word)...")
            chunk_count = processor.process_docx(file_path)
            print(f"Added {chunk_count} chunks")
        elif file_path.suffix.lower() in [".eml", ".msg"]:
            print(f"Processing {file_path.name} (Email)...")
            chunk_count = processor.process_email(file_path)
            print(f"Added {chunk_count} chunks")

if __name__ == "__main__":
    main()
